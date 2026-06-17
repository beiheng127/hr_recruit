from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from lxml import etree

doc = Document('2. 需求分析规格说明书-13.docx')

# Find paragraph 52 (2.2 heading)
heading_para = doc.paragraphs[52]
heading_elem = heading_para._p

# Create a new paragraph for the image (centered)
# First, add a paragraph at the end
new_para = doc.add_paragraph()
new_para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

# Add the image
run = new_para.add_run()
run.add_picture('usecase.png', width=Inches(6.0))

# Now move this paragraph from the end to after heading_para
new_elem = new_para._p

# Insert after heading
heading_elem.addnext(new_elem)

print('Image inserted after 2.2 heading')
print('Saving document...')

doc.save('2. 需求分析规格说明书-13.docx')
print('Done: 2. 需求分析规格说明书-13.docx saved with use case diagram')
