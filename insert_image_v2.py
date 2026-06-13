from docx import Document
from docx.shared import Inches
from lxml import etree
from docx.oxml.ns import qn

doc = Document('2. 需求分析规格说明书-13.docx')

# Find paragraph 52 (2.2 heading)
heading_para = doc.paragraphs[52]
heading_elem = heading_para._p

# The image is currently at the end of the document (from previous failed attempt).
# Let's find it and move it to the right location.
# Actually, let's just delete the old image paragraph and create a new one at the right location.

# Strategy: Use lxml to create a new <w:p> element with the image
# This is complex. Instead, let's:
# 1. Find the placeholder paragraph (index 53)
# 2. Clear it and add the image to it

placeholder = doc.paragraphs[53]
print(f'Placeholder text: {placeholder.text}')

# Clear the placeholder paragraph
# Remove all runs
for run in list(placeholder.runs):
    p_elem = placeholder._p
    r_elem = run._r
    p_elem.remove(r_elem)

# Now add the image to this paragraph
run = placeholder.add_run()
run.add_picture('usecase.png', width=Inches(6.0))
placeholder.alignment = 1  # Center

print('Image inserted into placeholder paragraph (index 53)')
print('Saving document...')

doc.save('2. 需求分析规格说明书-13.docx')
print('Done!')
